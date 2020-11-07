import os
import sys
import re
import time
from datetime import datetime
from docx import Document
import comtypes.client
from PyQt5 import QtCore, QtWidgets, QtGui
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication, QDialog, QPushButton, QWidget, QMainWindow
from PyQt5.uic import loadUi



global client_name
global doa
global attorney_name
global user_name
global user_email
global user_fax
global user_phone
global output_path
global widget
global lor_pl
global lop_pl
global rr_pl
global doc

#Path to files

path = os.getcwd()
settings_path = (path + "/settings")
lop_pl = (path + "\placeholders\LoP.docx")
rr_pl = (path + "\placeholders\RR.docx")
lor_pl = (path + "\placeholders\LoR.docx")
user_fax_pl = re.compile("User Fax")
user_email_pl = re.compile("User Email")
current_balance_pl = re.compile("Current Balance")
requested_balance_pl = re.compile("Requested Balance")

os.chdir (path+r"\gui")


#define replace_string , docxtopdf, and set values

def DocxtoPDF(in_file, out_file):
    word = comtypes.client.CreateObject('Word.Application')
    d = word.Documents.Open(in_file)
    d.SaveAs(out_file, FileFormat=17)
    d.Close()
    word.Quit()




def replace_string(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_string(cell, regex, replace)

def tabFocus(self):
    self.textEdit.setTabChangesFocus(True)
    self.textEdit_2.setTabChangesFocus(True)
    self.textEdit_3.setTabChangesFocus(True)
    self.textEdit_4.setTabChangesFocus(True)
    self.textEdit_5.setTabChangesFocus(True)
    self.setTabOrder(self.textEdit, self.textEdit_2)
    self.setTabOrder(self.textEdit_2, self.textEdit_3)
    self.setTabOrder(self.textEdit_3, self.textEdit_4)
    self.setTabOrder(self.textEdit_4, self.textEdit_5)

provider_int = 0

#define user settings

if os.path.exists(settings_path + "/user_name.txt"):
    with open(settings_path + '/user_name.txt', 'r') as user_name_file:
        user_name = user_name_file.read()

if os.path.exists(settings_path + "/user_email.txt"):
    with open(settings_path + '/user_email.txt', 'r') as user_email_file:
        user_email = user_email_file.read()

if os.path.exists(settings_path + "/user_phone.txt"):
    with open(settings_path + '/user_phone.txt', 'r') as user_phone_file:
        user_phone = user_phone_file.read()

if os.path.exists(settings_path + "/user_fax.txt"):
    with open(settings_path + '/user_fax.txt', 'r') as user_fax_file:
        user_fax = user_fax_file.read()

if os.path.exists("/output_path.txt"):
    with open('/output_path.txt', 'r') as output_path_file:
        output_path = output_path_file.read()


def retrievesettings(settingsfile):
    with open(settingsfile, "r") as settingsoutput:
        settingsoutput.read()


#placeholder values

#LoP and universal

provider_email_pl = re.compile("Provider Email")
provider_name_pl = re.compile("Provider Name")
client_name_pl = re.compile("Client Name")
doa_pl = re.compile("Date of Accident")
today_date_pl = re.compile("Today Date")
provider_address_pl = re.compile("Provider Address")
lop_amount_pl = re.compile("LOP Amount")
attorney_name_pl = re.compile("Attorney Name")

#LoR

def_insurance_pl = re.compile("Defendant Insurance")
claim_num_pl = re.compile("CNumber")
def_adjuster_name_pl = re.compile("AdjusterName")
def_adjuster_address_pl = re.compile("Defendant Adjuster Address")
def_adjuster_csz_pl = re.compile("Defendant Adjuster CSZ")
def_adjuster_fax_pl = re.compile("Defendant Adjuster Fax")
user_name_pl = re.compile("User Name")
user_number_pl = re.compile("User Number")

#RR

user_fax_pl = re.compile("User Fax")
user_email_pl = re.compile("User Email")
current_balance_pl = re.compile("Current Balance")
requested_balance_pl = re.compile("Requested Balance")

#Create date
now = datetime.now()
today_date = now.strftime('%m/%d/%Y')


#Create classes



#settings page

class settings(QDialog):
    def __init__(self):
        super(settings, self).__init__()
        loadUi("settings.ui", self)
        self.pushButton.clicked.connect(self.savesettings)
        self.pushButton_2.clicked.connect(self.movemain)
        self.setWindowTitle('Settings')
        tabFocus(self)

        if os.path.exists(settings_path + "/user_name.txt"):
            self.textEdit.insertPlainText(user_name)

        if os.path.exists(settings_path + "/user_email.txt"):
            self.textEdit_2.insertPlainText(user_email)

        if os.path.exists(settings_path + "/user_phone.txt"):
            self.textEdit_3.insertPlainText(user_phone)

        if os.path.exists(settings_path + "/user_fax.txt"):
            self.textEdit_4.insertPlainText(user_fax)

        if os.path.exists(settings_path + "/output_path.txt"):
            self.textEdit_5.insertPlainText(output_path)
        else:
            self.textEdit_5.insertPlainText('C:/')

    def savesettings(self):

        user_name = self.textEdit.toPlainText()
        self.writesettings("user_name", user_name)

        user_email = self.textEdit_2.toPlainText()
        self.writesettings("user_email", user_email)

        user_phone = self.textEdit_3.toPlainText()
        self.writesettings("user_phone", user_phone)

        user_fax = self.textEdit_4.toPlainText()
        self.writesettings("user_fax", user_fax)


    def writesettings(self, settingname, settingstr):
        with open(settings_path + "/" + settingname + ".txt", "w") as settingfile:
            settingfile.write(settingstr)




    def movemain(self):
        self.close()


#LOP page

class LOPpage(QDialog):
    def __init__(self):
        super(LOPpage, self).__init__()
        loadUi("LOP.ui", self)
        self.pushButton.clicked.connect(self.create_click)
        self.pushButton_2.clicked.connect(self.hide)
        self.setWindowTitle('LOP')

    def create_click(self):
        client_name = self.textEdit.toPlainText()
        doa = self.textEdit_2.toPlainText()
        attorney_name = self.textEdit_3.toPlainText()
        print(client_name + " "  + doa + " " + attorney_name)



#Reduction request page

class RRpage(QDialog):
    def __init__(self):
        super(RRpage, self).__init__()
        loadUi("RR.ui", self)
        self.pushButton.clicked.connect(self.create_click)
        self.pushButton_2.clicked.connect(self.hide)
        self.label_12.hide()
        self.label_13.hide()
        self.setWindowTitle('Reduction Request')

    def create_click(self):
        client_name = self.textEdit.toPlainText()
        doa = self.textEdit_2.toPlainText()
        attorney_name = self.textEdit_3.toPlainText()
        provider_name = self.textEdit_4.toPlainText()
        provider_email = self.textEdit_5.toPlainText()
        provider_address = self.textEdit_6.toPlainText()
        current_balance = self.textEdit_7.toPlainText()
        requested_balance = self.textEdit_8.toPlainText()
        doc = Document(rr_pl)

        if not os.path.exists('C:/DocuLegal/Reduction Requests/Word'):
            os.makedirs('C:/DocuLegal/Reduction Requests/Word')
        if not os.path.exists('C:/DocuLegal/Reduction Requests/PDF'):
            os.makedirs('C:/DocuLegal/Reduction Requests/PDF')

        replace_string(doc, client_name_pl, client_name)
        replace_string(doc, doa_pl, doa)
        replace_string(doc, attorney_name_pl, attorney_name)
        replace_string(doc, provider_name_pl, provider_name)
        replace_string(doc, provider_email_pl, provider_email)
        replace_string(doc, provider_address_pl, provider_address)
        replace_string(doc, current_balance_pl, current_balance)
        replace_string(doc, requested_balance_pl, requested_balance)
        replace_string(doc, user_name_pl, user_name)
        replace_string(doc, user_number_pl, user_phone)
        replace_string(doc, today_date_pl, today_date)
        print('check')
        doc.save('C:/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' REDUCTION REQUEST ' + provider_name.upper() + ".docx")
        print('check')
        if str(self.comboBox.currentText()) == "Docx and PDF":
            DocxtoPDF('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' REDUCTION REQUEST ' + provider_name.upper() + ".docx",
                      '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' REDUCTION REQUEST ' + provider_name.upper() + ".pdf")
            self.label_13.show()

        else:
            self.label_12.show()




#LOR page

class LORpage(QDialog):
    def __init__(self):
        super(LORpage, self).__init__()
        loadUi("LOR.ui", self)
        self.pushButton.clicked.connect(self.create_click)
        self.pushButton_2.clicked.connect(self.hide)
        self.setWindowTitle('LOR')

    def create_click(self):
        client_name = self.textEdit.toPlainText()
        doa = self.textEdit_2.toPlainText()
        attorney_name = self.textEdit_3.toPlainText()
        claim_num = self.textEdit_4.toPlainText()
        def_insurance = self.textEdit_5.toPlainText()
        def_adjuster_name = self.textEdit_6.toPlainText()
        def_adjuster_address = self.textEdit_7.toPlainText()
        def_adjuster_csz = self.textEdit_8.toPlainText()
        def_adjuster_fax = self.textEdit_9.toPlainText()

        doc = Document(lor_pl)
        if not os.path.exists('C:/DocuLegal/LORs/Word'):
            os.makedirs('C:/DocuLegal/LORs/Word')

        if not os.path.exists('C:/DocuLegal/LORs/PDF'):
            os.makedirs('C:/DocuLegal/LORs/PDF')

        replace_string(doc, client_name_pl, client_name)
        replace_string(doc, doa_pl, doa)
        replace_string(doc, def_insurance_pl, def_insurance)
        replace_string(doc, claim_num_pl, claim_num)
        replace_string(doc, def_adjuster_name_pl, def_adjuster_name)
        replace_string(doc, def_adjuster_address_pl, def_adjuster_address)
        replace_string(doc, def_adjuster_csz_pl, def_adjuster_csz)
        replace_string(doc, def_adjuster_fax_pl, def_adjuster_fax)
        replace_string(doc, user_name_pl, user_name)
        replace_string(doc, user_number_pl, user_phone)
        replace_string(doc, attorney_name_pl, attorney_name)
        replace_string(doc, def_adjuster_name_pl, def_adjuster_name)
        replace_string(doc, today_date_pl, today_date)


        doc.save('C:/DocuLegal/LORs/Word/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".docx")

        if str(self.comboBox.currentText()) == "Docx and PDF":

            DocxtoPDF('C:/DocuLegal/LORs/Word/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".docx",
                      'C:/DocuLegal/LORs/PDF/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".pdf")


# MRBR page

class MRBRpage(QDialog):
    def __init__(self):
        super(MRBRpage, self).__init__()
        loadUi("MRBR.ui", self)
        self.pushButton.clicked.connect(self.create_click)
        self.pushButton_2.clicked.connect(self.hide)
        self.setWindowTitle('MRBR')


# MRBR Hospital Page

class MRBRHpage(QDialog):
    def __init__(self):
        super(MRBRHpage, self).__init__()
        loadUi("MRBRH.ui", self)
        self.pushButton.clicked.connect(self.create_click)
        self.pushButton_2.clicked.connect(self.hide)
        self.setWindowTitle('MRBR - Hospital')

#Main page

class mainpage(QDialog):

    def __init__(self):
        super(mainpage, self).__init__()
        loadUi('mainpage.ui', self)
        self.pushButton.clicked.connect(self.moveLOP)
        self.pushButton_2.clicked.connect(self.moveLOR)
        self.pushButton_3.clicked.connect(self.moveRR)
        self.pushButton_4.clicked.connect(self.movesettings)
        self.settings = settings()
        self.LORpage = LORpage()
        self.LOPpage = LOPpage()
        self.RRpage = RRpage()
        #self.MRBRpage = MRBRpage()
        #self.MRBRHpage = MRBRHpage()
        self.LORpage.pushButton_2.clicked.connect(self.show)
        self.LOPpage.pushButton_2.clicked.connect(self.show)
        self.RRpage.pushButton_2.clicked.connect(self.show)
        #self.MRBRpage.pushButton_2.clicked.connect(self.show)
        #self.MRBRHpage.pushButton_2.clicked.connect(self.show)
        self.setWindowTitle('DocuLegal')

    def movesettings(self):
        self.settings.show()

    def moveLOR(self):
        self.LORpage.show()
        self.hide()
        global doc_type
        doc_type = "LOR"

    def moveLOP(self):
        self.LOPpage.show()
        self.hide()
        global doc_type
        doc_type = "LOP"

    def moveRR(self):
        self.RRpage.show()
        self.hide()
        global doc_type
        doc_type = "RR"




if os.path.exists(settings_path + "/user_name.txt") and os.path.exists(settings_path + "/user_email.txt") and os.path.exists(settings_path + "/user_phone.txt") and os.path.exists(settings_path + "/user_email.txt"):
    app = QApplication(sys.argv)
    global widget
    widget = mainpage()
    widget.show()
    sys.exit(app.exec_())

else:
    app = QApplication(sys.argv)
    widget = settings()
    widget.show()
    sys.exit(app.exec())
