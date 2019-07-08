import os
import re
import time
from datetime import datetime
from docx import Document
import comtypes.client

#Path to files
#new comment
#another comment
path = os.getcwd()
lop_pl = (path + "\placeholders\LoP.docx")
rr_pl = (path + "\placeholders\RR.docx")
lor_pl = (path + "\placeholders\LoR.docx")

#define replace_string , docxtopdf, and set values

def DocxtoPDF(inputFileName, outputFileName, formatType = 17):
    Word = comtypes.client.CreateObject("Word.Application")
    Word.Visible = 1

    if outputFileName[-3:] != 'pdf':
        outputFileName = outputFileName + ".pdf"
    deck = Word.Documents.Open(inputFileName)
    deck.SaveAs(outputFileName, formatType) # formatType = 32 for ppt to pdf
    deck.Close()
    Word.Quit()



def replace_string(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_string(cell, regex, replace)

provider_int = 0

#Placeholder values

#LoP and universal

provider_email_PL = re.compile("Provider Email")
provider_name_PL = re.compile("Provider Name")
client_name_PL = re.compile("Client Name")
doa_PL = re.compile("Date of Accident")
today_date_PL = re.compile("Today Date")
provider_address_PL = re.compile("Provider Address")
lop_amount_PL = re.compile("LOP Amount")
attorney_name_PL = re.compile("Attorney Name")

#LoR

def_insurance_pl = re.compile("Defendant Insurance")
claim_num_pl = re.compile("CNumber")
def_adjuster_name_pl = re.compile("AdjusterName")
def_adjuster_address_pl = re.compile("Defendant Adjuster Address")
def_adjuster_csz_pl = re.compile("Defendant Adjuster CSZ")
def_adjuster_fax_pl = re.compile("Defendant Adjuster Fax")
user_name_pl = re.compile("User Name")
user_number_pl = re.compile("User Number")

#RR

user_fax_pl = re.compile("User Fax")
user_email_pl = re.compile("User Email")
current_balance_pl = re.compile("Current Balance")
requested_balance_pl = re.compile("Requested Balance")

#Create date
now = datetime.now()
today_date = now.strftime('%m/%d/%Y')

#prompt users for input
doc_type = input('What Type of Document are you making? (RR/LOP/LOR): ')
client_name = input("Client Name: ")
doa = input('Date of Accident: ')
attorney_name= input("Attorney's Name: ")

if doc_type == "LOR":
    user_name = input("Your Name: ")
    user_num = input("Your Phone Number: ")
    def_insurance = input("Defendant's Insurance Company: ")
    claim_num = input("Claim Number: ")
    def_adjuster_name = input("Defendant Adjuster's Name: ")
    def_adjuster_address = input("Defendant Adjuster's Address: ")
    def_adjuster_csz = input("Defendant Adjuster's City, State, Zip: ")
    def_adjuster_fax = input("Defendant Adjuster's Fax Number: ")


if doc_type == "RR":
    user_name = input("Your Name: ")
    user_email = input("Your Email: ")
    user_fax = input("Your Fax Number: ")

if doc_type == "LOP":
    lop_amount = input("LoP Amount: ")


if not doc_type == "LOR":
    provider_num = input("How many providers?(1-10): ")
    provider_int = int(provider_num)
    provider_name = input("First Provider Name: ")
    provider_email = input(provider_name + "'s Email/Fax: ")
    provider_address = input(provider_name + "'s Address: ")

if doc_type == "RR":
    current_balance = input(provider_name + "'s Current Balance: ")
    requested_balance = input(provider_name + "'s Requested Balance: ")






if provider_int > 1:
    provider_name2 = input("Second Provider Name: ")
    provider_email2 = input(provider_name2 + "'s Email/Fax: ")
    provider_address2 = input(provider_name2 + "'s Address: ")
    if doc_type == "RR":
        current_balance2 = input(provider_name2 + "'s Current Balance: ")
        requested_balance2 = input(provider_name2 + "'s Requested Balance: ")

if provider_int > 2:
    provider_name3 = input("Third Provider Name: ")
    provider_email3 = input(provider_name3 + "'s Email/Fax: ")
    provider_address3 = input(provider_name3 + "'s Address: ")

    if doc_type == "RR":
        current_balance3 = input(provider_name3 + "'s Current Balance: ")
        requested_balance3 = input(provider_name3 + "'s Requested Balance: ")

if provider_int > 3:
    provider_name4 = input("Fourth Provider Name: ")
    provider_email4 = input(provider_name4 + "'s Email/Fax: ")
    provider_address4 = input(provider_name4 + "'s Address: ")

    if doc_type == "RR":
        current_balance4 = input(provider_name4 + "'s Current Balance: ")
        requested_balance4 = input(provider_name4 + "'s Requested Balance: ")

if provider_int > 4:
    provider_name5 = input("Fifth Provider Name: ")
    provider_email5 = input(provider_name5 + "'s Email/Fax: ")
    provider_address5 = input(provider_name5 + "'s Address: ")

    if doc_type == "RR":
        current_balance5 = input(provider_name5 + "'s Current Balance: ")
        requested_balance5 = input(provider_name5 + "'s Requested Balance: ")

if provider_int > 5:
    provider_name6 = input("Sixth Provider Name: ")
    provider_email6 = input(provider_name6 + "'s Email/Fax: ")
    provider_address6 = input(provider_name6 + "'s Address: ")

    if doc_type == "RR":
        current_balance6 = input(provider_name6 + "'s Current Balance: ")
        requested_balance6 = input(provider_name6 + "'s Requested Balance: ")

if provider_int > 6:
    provider_name7 = input("Seventh Provider Name: ")
    provider_email7 = input(provider_name7 + "'s Email/Fax: ")
    provider_address7 = input(provider_name7 + "'s Address: ")

    if doc_type == "RR":
        current_balance7 = input(provider_name7 + "'s Current Balance: ")
        requested_balance7 = input(provider_name7 + "'s Requested Balance: ")

if provider_int > 7:
    provider_name8 = input("Eighth Provider Name: ")
    provider_email8 = input(provider_name8 + "'s Email/Fax: ")
    provider_address8 = input(provider_name8 + "'s Address: ")

    if doc_type == "RR":
        current_balance8 = input(provider_name8 + "'s Current Balance: ")
        requested_balance8 = input(provider_name8 + "'s Requested Balance: ")

if provider_int > 8:
    provider_name9 = input("Ninth Provider Name: ")
    provider_email9 = input(provider_name9 + "'s Email/Fax: ")
    provider_address9 = input(provider_name9 + "'s Address: ")

    if doc_type == "RR":
        current_balance9 = input(provider_name9 + "'s Current Balance: ")
        requested_balance9 = input(provider_name9 + "'s Requested Balance: ")

if provider_int > 9:
    provider_name10 = input("Tenth Provider Name: ")
    provider_email10 = input(provider_name10 + "'s Email/Fax: ")
    provider_address10 = input(provider_name10 + "'s Address: ")

    if doc_type == "RR":
        current_balance10 = input(provider_name10 + "'s Current Balance: ")
        requested_balance10 = input(provider_name10 + "'s Requested Balance: ")


#link placeholder document and make path
if doc_type == "LOP":

    doc = Document(lop_pl)

    if not os.path.exists('/DocuLegal/LOPs/Word'):
        os.makedirs('/DocuLegal/LOPs/Word')

    if not os.path.exists('/DocuLegal/LOPs/PDF'):
        os.makedirs('/DocuLegal/LOPs/PDF')

    doc.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name.upper() + ".docx")


if doc_type == "RR":

    doc = Document(rr_pl)

    if not os.path.exists('/DocuLegal/Reduction Requests/Word'):
        os.makedirs('/DocuLegal/Reduction Requests/Word')

    if not os.path.exists('/DocuLegal/Reduction Requests/PDF'):
        os.makedirs('/DocuLegal/Reduction Requests/PDF')

    doc.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name.upper() + ".docx")


if doc_type == "LOR":

    doc = Document(lor_pl)

    if not os.path.exists('/DocuLegal/LORs/Word'):
        os.makedirs('/DocuLegal/LORs/Word')

    if not os.path.exists('/DocuLegal/LORs/PDF'):
        os.makedirs('/DocuLegal/LORs/PDF')

    doc.save('/DocuLegal/LORs/Word/' + client_name.upper() + ' LOR ' + def_insurance.upper()+".docx")

#replace placeholders with user input
replace_string(doc, client_name_PL, client_name)
replace_string(doc, doa_PL, doa)
replace_string(doc, today_date_PL, today_date)


#LOR
if doc_type == "LOR":
    replace_string(doc, client_name_PL, client_name)
    replace_string(doc, def_insurance_pl, def_insurance)
    replace_string(doc, claim_num_pl, claim_num)
    replace_string(doc, def_adjuster_name_pl, def_adjuster_name)
    replace_string(doc, def_adjuster_address_pl, def_adjuster_address)
    replace_string(doc, def_adjuster_csz_pl, def_adjuster_csz)
    replace_string(doc, def_adjuster_fax_pl, def_adjuster_fax)
    replace_string(doc, user_name_pl, user_name)
    replace_string(doc, user_number_pl, user_num)
    replace_string(doc, attorney_name_PL, attorney_name)
    replace_string(doc, def_adjuster_name_pl, def_adjuster_name)

#LOP
if doc_type == "LOP":
    replace_string(doc, provider_address_PL, provider_address)
    replace_string(doc, provider_email_PL, provider_email)
    replace_string(doc, provider_name_PL, provider_name)
    replace_string(doc, lop_amount_PL, lop_amount)
    replace_string(doc, attorney_name_PL, attorney_name)

#RR
if doc_type == "RR":
    replace_string(doc, user_name_pl, user_name)
    replace_string(doc, user_fax_pl, user_fax)
    replace_string(doc, user_email_pl, user_email)

    replace_string(doc, current_balance_pl, current_balance)
    replace_string(doc, requested_balance_pl, requested_balance)
    replace_string(doc, provider_email_PL, provider_email)
    replace_string(doc, provider_name_PL, provider_name)

#save
if doc_type == "LOP":
    doc.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name.upper() + ".docx")
if doc_type == "RR":
    doc.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name.upper() + ".docx")
if doc_type == "LOR":
    doc.save('/DocuLegal/LORs/Word/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".docx")

#loop for extra providers
if provider_int > 1:
    if doc_type == "LOP":
        doc2 = Document(lop_pl)
        doc2.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name2.upper() + ".docx")
        replace_string(doc2, lop_amount_PL, lop_amount)
        replace_string(doc2, provider_address_PL, provider_address2)
        replace_string(doc2, provider_email_PL, provider_email2)
        replace_string(doc2, provider_name_PL, provider_name2)

    if doc_type == "RR":
        doc2 = Document(rr_pl)
        doc2.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name2.upper() + ".docx")
        replace_string(doc2, provider_address_PL, provider_address2)
        replace_string(doc2, provider_email_PL, provider_email2)
        replace_string(doc2, provider_name_PL, provider_name2)
        replace_string(doc2, user_name_pl, user_name)
        replace_string(doc2, user_fax_pl, user_fax)
        replace_string(doc2, user_email_pl, user_email)
        replace_string(doc2, current_balance_pl, current_balance2)
        replace_string(doc2, requested_balance_pl, requested_balance2)

    replace_string(doc2, client_name_PL, client_name)
    replace_string(doc2, doa_PL, doa)
    replace_string(doc2, today_date_PL, today_date)
    replace_string(doc2, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc2.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name2.upper() + ".docx")
    if doc_type == "RR":
        doc2.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name2.upper() + ".docx")


if provider_int > 2:

    if doc_type == "LOP":
        doc3 = Document(lop_pl)
        doc3.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name3.upper() + ".docx")
        replace_string(doc3, lop_amount_PL, lop_amount)
        replace_string(doc3, provider_address_PL, provider_address3)
        replace_string(doc3, provider_email_PL, provider_email3)
        replace_string(doc3, provider_name_PL, provider_name3)

    if doc_type == "RR":
        doc3 = Document(rr_pl)
        doc3.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name3.upper() + ".docx")
        replace_string(doc3, provider_address_PL, provider_address3)
        replace_string(doc3, provider_email_PL, provider_email3)
        replace_string(doc3, provider_name_PL, provider_name3)
        replace_string(doc3, user_name_pl, user_name)
        replace_string(doc3, user_fax_pl, user_fax)
        replace_string(doc3, user_email_pl, user_email)
        replace_string(doc3, current_balance_pl, current_balance3)
        replace_string(doc3, requested_balance_pl, requested_balance3)

    replace_string(doc3, client_name_PL, client_name)
    replace_string(doc3, doa_PL, doa)
    replace_string(doc3, today_date_PL, today_date)
    replace_string(doc3, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc3.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name3.upper() + ".docx")
    if doc_type == "RR":
        doc3.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name3.upper() + ".docx")

if provider_int > 3:

    if doc_type == "LOP":
        doc4 = Document(lop_pl)
        doc4.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name4.upper() + ".docx")
        replace_string(doc4, lop_amount_PL, lop_amount)
        replace_string(doc4, provider_address_PL, provider_address4)
        replace_string(doc4, provider_email_PL, provider_email4)
        replace_string(doc4, provider_name_PL, provider_name4)

    if doc_type == "RR":
        doc4 = Document(rr_pl)
        doc4.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name4.upper() + ".docx")
        replace_string(doc4, provider_address_PL, provider_address4)
        replace_string(doc4, provider_email_PL, provider_email4)
        replace_string(doc4, provider_name_PL, provider_name4)
        replace_string(doc4, user_name_pl, user_name)
        replace_string(doc4, user_fax_pl, user_fax)
        replace_string(doc4, user_email_pl, user_email)
        replace_string(doc4, current_balance_pl, current_balance4)
        replace_string(doc4, requested_balance_pl, requested_balance4)

    replace_string(doc4, client_name_PL, client_name)
    replace_string(doc4, doa_PL, doa)
    replace_string(doc4, today_date_PL, today_date)
    replace_string(doc4, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc4.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name4.upper() + ".docx")
    if doc_type == "RR":
        doc4.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name4.upper() + ".docx")


if provider_int > 4:

    if doc_type == "LOP":
        doc5 = Document(lop_pl)
        doc5.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name5.upper() + ".docx")
        replace_string(doc5, lop_amount_PL, lop_amount)
        replace_string(doc5, provider_address_PL, provider_address5)
        replace_string(doc5, provider_email_PL, provider_email5)
        replace_string(doc5, provider_name_PL, provider_name5)

    if doc_type == "RR":
        doc5 = Document(rr_pl)
        doc5.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name5.upper() + ".docx")
        replace_string(doc5, provider_address_PL, provider_address5)
        replace_string(doc5, provider_email_PL, provider_email5)
        replace_string(doc5, provider_name_PL, provider_name5)
        replace_string(doc5, user_name_pl, user_name)
        replace_string(doc5, user_fax_pl, user_fax)
        replace_string(doc5, user_email_pl, user_email)
        replace_string(doc5, current_balance_pl, current_balance5)
        replace_string(doc5, requested_balance_pl, requested_balance5)

    replace_string(doc5, client_name_PL, client_name)
    replace_string(doc5, doa_PL, doa)
    replace_string(doc5, today_date_PL, today_date)
    replace_string(doc5, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc5.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name5.upper() + ".docx")
    if doc_type == "RR":
        doc5.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name5.upper() + ".docx")


if provider_int > 5:

    if doc_type == "LOP":
        doc6 = Document(lop_pl)
        doc6.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name6.upper() + ".docx")
        replace_string(doc6, lop_amount_PL, lop_amount)
        replace_string(doc6, provider_address_PL, provider_address6)
        replace_string(doc6, provider_email_PL, provider_email6)
        replace_string(doc6, provider_name_PL, provider_name6)

    if doc_type == "RR":
        doc6 = Document(rr_pl)
        doc6.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name6.upper() + ".docx")
        replace_string(doc6, provider_address_PL, provider_address6)
        replace_string(doc6, provider_email_PL, provider_email6)
        replace_string(doc6, provider_name_PL, provider_name6)
        replace_string(doc6, user_name_pl, user_name)
        replace_string(doc6, user_fax_pl, user_fax)
        replace_string(doc6, user_email_pl, user_email)
        replace_string(doc6, current_balance_pl, current_balance6)
        replace_string(doc6, requested_balance_pl, requested_balance6)

    replace_string(doc6, client_name_PL, client_name)
    replace_string(doc6, doa_PL, doa)
    replace_string(doc6, today_date_PL, today_date)
    replace_string(doc6, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc6.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name6.upper() + ".docx")
    if doc_type == "RR":
        doc6.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name6.upper() + ".docx")


if provider_int > 6:

    if doc_type == "LOP":
        doc7 = Document(lop_pl)
        doc7.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name7.upper() + ".docx")
        replace_string(doc7, lop_amount_PL, lop_amount)
        replace_string(doc7, provider_address_PL, provider_address7)
        replace_string(doc7, provider_email_PL, provider_email7)
        replace_string(doc7, provider_name_PL, provider_name7)

    if doc_type == "RR":
        doc7 = Document(rr_pl)
        doc7.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name7.upper() + ".docx")
        replace_string(doc7, provider_address_PL, provider_address7)
        replace_string(doc7, provider_email_PL, provider_email7)
        replace_string(doc7, provider_name_PL, provider_name7)
        replace_string(doc7, user_name_pl, user_name)
        replace_string(doc7, user_fax_pl, user_fax)
        replace_string(doc7, user_email_pl, user_email)
        replace_string(doc7, current_balance_pl, current_balance7)
        replace_string(doc7, requested_balance_pl, requested_balance7)

    replace_string(doc7, client_name_PL, client_name)
    replace_string(doc7, doa_PL, doa)
    replace_string(doc7, today_date_PL, today_date)
    replace_string(doc7, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc7.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name7.upper() + ".docx")
    if doc_type == "RR":
        doc7.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name7.upper() + ".docx")


if provider_int > 7:

    if doc_type == "LOP":
        doc8 = Document(lop_pl)
        doc8.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name8.upper() + ".docx")
        replace_string(doc8, lop_amount_PL, lop_amount)
        replace_string(doc8, provider_address_PL, provider_address8)
        replace_string(doc8, provider_email_PL, provider_email8)
        replace_string(doc8, provider_name_PL, provider_name8)

    if doc_type == "RR":
        doc8 = Document(rr_pl)
        doc8.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name8.upper() + ".docx")
        replace_string(doc8, provider_address_PL, provider_address8)
        replace_string(doc8, provider_email_PL, provider_email8)
        replace_string(doc8, provider_name_PL, provider_name8)
        replace_string(doc8, user_name_pl, user_name)
        replace_string(doc8, user_fax_pl, user_fax)
        replace_string(doc8, user_email_pl, user_email)
        replace_string(doc8, current_balance_pl, current_balance8)
        replace_string(doc8, requested_balance_pl, requested_balance8)

    replace_string(doc8, client_name_PL, client_name)
    replace_string(doc8, doa_PL, doa)
    replace_string(doc8, today_date_PL, today_date)
    replace_string(doc8, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc8.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name8.upper() + ".docx")
    if doc_type == "RR":
        doc8.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name8.upper() + ".docx")


if provider_int > 8:

    if doc_type == "LOP":
        doc9 = Document(lop_pl)
        doc9.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name9.upper() + ".docx")
        replace_string(doc9, lop_amount_PL, lop_amount)
        replace_string(doc9, provider_address_PL, provider_address9)
        replace_string(doc9, provider_email_PL, provider_email9)
        replace_string(doc9, provider_name_PL, provider_name9)

    if doc_type == "RR":
        doc9 = Document(rr_pl)
        doc9.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name9.upper() + ".docx")
        replace_string(doc9, provider_address_PL, provider_address9)
        replace_string(doc9, provider_email_PL, provider_email9)
        replace_string(doc9, provider_name_PL, provider_name9)
        replace_string(doc9, user_name_pl, user_name)
        replace_string(doc9, user_fax_pl, user_fax)
        replace_string(doc9, user_email_pl, user_email)
        replace_string(doc9, current_balance_pl, current_balance9)
        replace_string(doc9, requested_balance_pl, requested_balance9)

    replace_string(doc9, client_name_PL, client_name)
    replace_string(doc9, doa_PL, doa)
    replace_string(doc9, today_date_PL, today_date)
    replace_string(doc9, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc9.save('/DocuLegal/LOPs/' + client_name.upper() + ' LOP ' + provider_name9.upper() + ".docx")
    if doc_type == "RR":
        doc9.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name9.upper() + ".docx")


if provider_int > 9:

    if doc_type == "LOP":
        doc10 = Document(lop_pl)
        doc10.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name10.upper() + ".docx")
        replace_string(doc10, lop_amount_PL, lop_amount)
        replace_string(doc10, provider_address_PL, provider_address10)
        replace_string(doc10, provider_email_PL, provider_email10)
        replace_string(doc10, provider_name_PL, provider_name10)

    if doc_type == "RR":
        doc10 = Document(rr_pl)
        doc10.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name10.upper() + ".docx")
        replace_string(doc10, provider_address_PL, provider_address10)
        replace_string(doc10, provider_email_PL, provider_email10)
        replace_string(doc10, provider_name_PL, provider_name10)
        replace_string(doc10, user_name_pl, user_name)
        replace_string(doc10, user_fax_pl, user_fax)
        replace_string(doc10, user_email_pl, user_email)
        replace_string(doc10, current_balance_pl, current_balance10)
        replace_string(doc10, requested_balance_pl, requested_balance10)

    replace_string(doc10, client_name_PL, client_name)
    replace_string(doc10, doa_PL, doa)
    replace_string(doc10, today_date_PL, today_date)
    replace_string(doc10, attorney_name_PL, attorney_name)


    if doc_type == "LOP":
        doc10.save('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name10.upper() + ".docx")
    if doc_type == "RR":
        doc10.save('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name10.upper() + ".docx")



#notify user of completion
if provider_int == 1:
    print('Document sent to "DocuLegal" folder')
else:
    print('Documents sent to "DocuLegal" folder')

MakePDF = input("Would you also like PDF versions of these files? (Y/N): ")

if MakePDF == "Y":
    print("Please Wait...")
    if doc_type == "LOR":
        DocxtoPDF('/DocuLegal/LORs/Word/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".docx",
                  '/DocuLegal/LORs/PDF/' + client_name.upper() + ' LOR ' + def_insurance.upper() + ".pdf")


    if doc_type == "LOP":

        DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name.upper() + ".docx",
                  '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name.upper() + ".pdf")

        if provider_int > 1:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name2.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name2.upper() + ".pdf")

        if provider_int > 2:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name3.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name3.upper() + ".pdf")

        if provider_int > 3:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name4.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name4.upper() + ".pdf")

        if provider_int > 4:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name5.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name5.upper() + ".pdf")

        if provider_int > 5:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name6.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name6.upper() + ".pdf")

        if provider_int > 6:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name7.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name7.upper() + ".pdf")

        if provider_int > 7:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name8.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name8.upper() + ".pdf")

        if provider_int > 8:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name9.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name9.upper() + ".pdf")

        if provider_int > 9:
            DocxtoPDF('/DocuLegal/LOPs/Word/' + client_name.upper() + ' LOP ' + provider_name10.upper() + ".docx",
                      '/DocuLegal/LOPs/PDF/' + client_name.upper() + ' LOP ' + provider_name10.upper() + ".pdf")
    if doc_type == "RR":

        DocxtoPDF('/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name.upper() + ".docx",
                  '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name.upper() + ".pdf")

        if provider_int > 1:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name2.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name2.upper() + ".pdf")

        if provider_int > 2:
                DocxtoPDF(
                    '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name3.upper() + ".docx",
                    '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name3.upper() + ".pdf")

        if provider_int > 3:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name4.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name4.upper() + ".pdf")

        if provider_int > 4:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name5.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name5.upper() + ".pdf")

        if provider_int > 5:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name6.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name6.upper() + ".pdf")

        if provider_int > 6:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name7.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name7.upper() + ".pdf")

        if provider_int > 7:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name8.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name8.upper() + ".pdf")

        if provider_int > 8:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name9.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name9.upper() + ".pdf")

        if provider_int > 9:
            DocxtoPDF(
                '/DocuLegal/Reduction Requests/Word/' + client_name.upper() + ' Reduction Request ' + provider_name10.upper() + ".docx",
                '/DocuLegal/Reduction Requests/PDF/' + client_name.upper() + ' Reduction Request ' + provider_name10.upper() + ".pdf")

    if provider_int == 1:
        print('Document sent to "DocuLegal" folder')
    else:
        print('Documents sent to "DocuLegal" folder')

time.sleep(5)